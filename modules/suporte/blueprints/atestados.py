"""Rotas de atestados (importadas do sistema Viva Rio)."""
from __future__ import annotations

from datetime import datetime
import calendar
import io
import os
import tempfile
import threading
import uuid
import unicodedata
import smtplib
import subprocess
from email.message import EmailMessage
from email.utils import formataddr

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import (
    Blueprint,
    current_app,
    flash,
    jsonify,
    make_response,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from utils.helpers import (
    wants_json as _wants_json,
    submit_bg_task,
)
from flask_login import current_user, login_required
from sqlalchemy.orm import joinedload

from extensions import db
from modules.propostas.blueprints.auth.permissions_utils import normalize_role_key, raw_permissions, current_permissions
from modules.audit.utils import write_audit_external
from ..models import (
    AtestadoArquivo,
    AtestadoEmail,
    AtestadoLogEnvio,
    AtestadoTask,
)


atestados_bp = Blueprint("atestados_bp", __name__, url_prefix="/assistencia/atestados")

MESES_MAP = {
    "JANEIRO": 1,
    "FEVEREIRO": 2,
    "MARÇO": 3,
    "ABRIL": 4,
    "MAIO": 5,
    "JUNHO": 6,
    "JULHO": 7,
    "AGOSTO": 8,
    "SETEMBRO": 9,
    "OUTUBRO": 10,
    "NOVEMBRO": 11,
    "DEZEMBRO": 12,
}


def _dept_names(user=None) -> set[str]:
    actor = user or current_user
    names: set[str] = set()
    try:
        for name in getattr(actor, "department_names", []) or []:
            cleaned = (name or "").strip()
            if cleaned:
                normalized = unicodedata.normalize("NFKD", cleaned)
                normalized = "".join(ch for ch in normalized if not unicodedata.combining(ch))
                names.add(normalized.upper())
    except Exception:
        return set()
    return names


def _role_key() -> str:
    return normalize_role_key(getattr(current_user, "tipo", None) or session.get("tipo"))


def _has_assist_admin_permission() -> bool:
    perms = current_permissions()
    return perms.get("admin_assistencia") or perms.get("admin_suporte")




def _deny_access(area_label: str):
    if _wants_json():
        return jsonify({"ok": False, "message": "Você não tem permissão para acessar esta área."}), 403
    flash(
        "Você não tem permissão para acessar esta área. Procure seu superior caso precise de acesso.",
        "warning",
    )
    return redirect(url_for("sem_permissao", area=area_label))


def _nav_counts() -> dict[str, int]:
    return {
        "anexos": AtestadoArquivo.query.count(),
        "historico": AtestadoLogEnvio.query.count(),
        "envios": AtestadoEmail.query.count(),
    }


@atestados_bp.before_request
def _check_permissions():
    from flask import request
    if "/api/" in getattr(request, "path", ""):
        return
    endpoint = getattr(request, "endpoint", "") or ""
    if endpoint and not endpoint.startswith("atestados_bp."):
        return
    if not current_user.is_authenticated:
        return
    allowed_depts = {"ASSISTENCIA TECNICA", "ESTOQUE", "OFICINA"}
    dept_names = _dept_names()

    role_key = _role_key()
    if role_key in ("admin", "gestor"):
        return
    perms = current_permissions()
    if _has_assist_admin_permission() or perms.get("assistencia_atestados") or (dept_names & allowed_depts):
        return
    return _deny_access("Assistencia tecnica")


def _set_highlight_color(run, color: str = "yellow") -> None:
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color)
    rpr = run._r.get_or_add_rPr()
    rpr.append(highlight)


def _converter_docx_para_pdf(docx_path: str, pdf_path: str) -> None:
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
        check=True,
    )


def _resolve_atestado_mail_settings() -> dict[str, object]:
    cfg = current_app.config
    defaults = {
        "host": "smtp.sollustecnologia.com",
        "port": 587,
        "username": "contratos.automatico@sollustecnologia.com",
        "password": None,  # Must be configured via ATESTADOS_SMTP_PASSWORD or MAIL_PASSWORD env var
        "from_email": "atestados@sollusgroup.com",
        "from_name": "Setor de Atestados",
        "reply_to": "atestados@sollusgroup.com",
        "cc": ["atestados@sollusgroup.com"],
        "use_tls": True,
    }
    return {
        "host": cfg.get("ATESTADOS_SMTP_HOST") or cfg.get("MAIL_SERVER") or defaults["host"],
        "port": int(cfg.get("ATESTADOS_SMTP_PORT") or cfg.get("MAIL_PORT") or defaults["port"]),
        "username": cfg.get("ATESTADOS_SMTP_USERNAME") or cfg.get("MAIL_USERNAME") or defaults["username"],
        "password": cfg.get("ATESTADOS_SMTP_PASSWORD") or cfg.get("MAIL_PASSWORD") or defaults["password"],
        "from_email": cfg.get("ATESTADOS_FROM_EMAIL") or defaults["from_email"],
        "from_name": cfg.get("ATESTADOS_FROM_NAME") or defaults["from_name"],
        "reply_to": cfg.get("ATESTADOS_REPLY_TO") or defaults["reply_to"],
        "cc": cfg.get("ATESTADOS_EMAIL_CC") or defaults["cc"],
        "use_tls": cfg.get("ATESTADOS_SMTP_USE_TLS") if cfg.get("ATESTADOS_SMTP_USE_TLS") is not None else cfg.get("MAIL_USE_TLS", defaults["use_tls"]),
        "use_ssl": cfg.get("ATESTADOS_SMTP_USE_SSL") if cfg.get("ATESTADOS_SMTP_USE_SSL") is not None else cfg.get("MAIL_USE_SSL", False),
    }


def _enviar_email_com_anexo(
    arquivo: AtestadoArquivo, pdf_file_path: str, mes_novo: str, ano_novo: str, data_limite: str
) -> tuple[str, str]:
    emails = [email.endereco for email in arquivo.emails]
    if not emails:
        try:
            write_audit_external(
                entity_type="atestados_email",
                entity_id=arquivo.id,
                action="email_skip",
                message="Envio de atestado ignorado: sem destinatarios.",
                after={"arquivo": arquivo.nome, "status": "no_email"},
            )
        except Exception:
            current_app.logger.exception("Falha ao auditar envio de atestado sem emails.")
        return ("no_email", "Nenhum email associado ao arquivo.")

    subject = os.path.splitext(arquivo.nome)[0]
    body = f"""
    <p>Prezados, bom dia!</p>
    <p>Segue em anexo o <b><span style="background-color: yellow">ATESTADO DE SERVIÇOS PRESTADOS</span></b> para validação referente ao <b><span style="background-color: cyan">MÊS DE {mes_novo}</span></b> do relógio de ponto biométrico.</p>
    <p style="color: red"><b>Favor enviar o atestado assinado até dia {data_limite}.</b></p>
    <p>Atenciosamente,</p>
    <p>Setor de Atestados<br>
    <a href="mailto:atestados@sollusgroup.com">atestados@sollusgroup.com</a><br>
    Av. Brasil, 31.904 - Bangu - Rio de Janeiro<br>
    21 2413-3203 | 22 2733-3722<br>
    27 3072-4863 | 41 3797-5093</p>
    """

    settings = _resolve_atestado_mail_settings()
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = formataddr((settings.get("from_name") or "", settings.get("from_email") or ""))
    if settings.get("reply_to"):
        msg["Reply-To"] = settings.get("reply_to")
    msg["To"] = ", ".join(emails)
    cc_list = settings.get("cc") or []
    if isinstance(cc_list, str):
        cc_list = [cc_list]
    if cc_list:
        msg["Cc"] = ", ".join(cc_list)
    msg.add_alternative(body, subtype="html")

    arquivo_pdf_nome = os.path.splitext(arquivo.nome)[0] + ".pdf"
    with open(pdf_file_path, "rb") as f:
        file_data = f.read()
        msg.add_attachment(file_data, maintype="application", subtype="pdf", filename=arquivo_pdf_nome)

    import ssl
    skip_tls_verify = str(current_app.config.get("MAIL_SKIP_TLS_VERIFY", "")).strip() == "1"
    context = ssl.create_default_context()
    if skip_tls_verify:
        context.check_hostname = False
        context.verify_mode = ssl.CERT_NONE

    try:
        use_ssl = settings.get("use_ssl", False) or settings["port"] == 465
        if use_ssl:
            smtp_client = smtplib.SMTP_SSL(settings["host"], settings["port"], context=context, timeout=30)
        else:
            smtp_client = smtplib.SMTP(settings["host"], settings["port"], timeout=30)

        with smtp_client as server:
            server.ehlo()
            if not use_ssl and settings.get("use_tls"):
                server.starttls(context=context)
                server.ehlo()
            if settings.get("username"):
                server.login(settings["username"], settings.get("password") or "")
            server.send_message(msg)
        try:
            write_audit_external(
                entity_type="atestados_email",
                entity_id=arquivo.id,
                action="email_send",
                message="Envio de atestado concluido.",
                after={
                    "arquivo": arquivo.nome,
                    "destinatários": emails,
                    "cc": cc_list,
                    "status": "success",
                },
            )
        except Exception:
            current_app.logger.exception("Falha ao auditar envio de atestado.")
        return ("success", "Sucesso")
    except Exception as exc:
        current_app.logger.exception("Falha ao enviar email do atestado: %s", exc)
        try:
            write_audit_external(
                entity_type="atestados_email",
                entity_id=arquivo.id,
                action="email_error",
                message="Falha ao enviar atestado.",
                after={
                    "arquivo": arquivo.nome,
                    "destinatários": emails,
                    "status": "error",
                    "erro": str(exc),
                },
            )
        except Exception:
            current_app.logger.exception("Falha ao auditar erro de envio de atestado.")
        return ("error", f"Erro: {exc}")


@atestados_bp.route("/", methods=["GET"])
@login_required
def index():
    page = request.args.get("page", 1, type=int)
    nome_filter = request.args.get("nome", "").strip()
    email_filter = request.args.get("email", "").strip()

    query = AtestadoArquivo.query
    if nome_filter:
        query = query.filter(AtestadoArquivo.nome.ilike(f"%{nome_filter}%"))
    if email_filter:
        query = (
            query.join(AtestadoArquivo.emails)
            .filter(AtestadoEmail.endereco.ilike(f"%{email_filter}%"))
            .distinct()
        )

    query = query.options(joinedload(AtestadoArquivo.emails))
    arquivos = query.order_by(AtestadoArquivo.id.asc()).paginate(page=page, per_page=10)

    stats = [
        {"label": "Arquivos", "value": AtestadoArquivo.query.count()},
        {"label": "Emails", "value": AtestadoEmail.query.count()},
    ]

    return render_template(
        "admin/assistencia/atestados/index.html",
        arquivos=arquivos,
        stats=stats,
        nome_filter=nome_filter,
        email_filter=email_filter,
        nav_counts=_nav_counts(),
        active_tab="anexos",
    )


@atestados_bp.route("/upload", methods=["GET", "POST"])
@login_required
def upload():
    if request.method == "POST":
        arquivo_enviado = request.files.get("arquivo")
        emails = (request.form.get("emails") or "").strip()

        if arquivo_enviado and emails:
            nome_original = arquivo_enviado.filename or ""
            if not nome_original.lower().endswith(".docx"):
                flash("Envie apenas arquivos DOCX.", "warning")
                return redirect(url_for("atestados_bp.upload"))

            conteudo = arquivo_enviado.read()
            arquivo_existente = AtestadoArquivo.query.filter_by(nome=nome_original).first()
            if arquivo_existente:
                flash("Um arquivo com esse nome já existe.", "warning")
                return redirect(url_for("atestados_bp.upload"))

            novo_arquivo = AtestadoArquivo(nome=nome_original, conteudo=conteudo)
            db.session.add(novo_arquivo)
            db.session.commit()

            lista_emails = [email.strip() for email in emails.split(",") if email.strip()]
            for email_endereco in lista_emails:
                db.session.add(AtestadoEmail(endereco=email_endereco, arquivo_id=novo_arquivo.id))

            db.session.commit()
            flash("Arquivo e emails salvos com sucesso!", "success")
            return redirect(url_for("atestados_bp.index"))

        flash("Por favor, envie um arquivo e insira emails.", "warning")

    return render_template("admin/assistencia/atestados/upload.html")


@atestados_bp.route("/<int:arquivo_id>/editar", methods=["GET", "POST"])
@login_required
def editar(arquivo_id: int):
    arquivo = AtestadoArquivo.query.get_or_404(arquivo_id)

    if request.method == "POST":
        arquivo_enviado = request.files.get("arquivo")
        if arquivo_enviado and arquivo_enviado.filename:
            nome_original = arquivo_enviado.filename
            if not nome_original.lower().endswith(".docx"):
                flash("Envie apenas arquivos DOCX.", "warning")
                return redirect(url_for("atestados_bp.editar", arquivo_id=arquivo.id))
            arquivo.nome = nome_original
            arquivo.conteudo = arquivo_enviado.read()

        novos_emails = (request.form.get("emails") or "").strip()
        if novos_emails:
            AtestadoEmail.query.filter_by(arquivo_id=arquivo.id).delete()
            lista_emails = [email.strip() for email in novos_emails.split(",") if email.strip()]
            for email_endereco in lista_emails:
                db.session.add(AtestadoEmail(endereco=email_endereco, arquivo_id=arquivo.id))

        db.session.commit()
        flash("Arquivo e emails atualizados com sucesso!", "success")
        return redirect(url_for("atestados_bp.index"))

    emails = ", ".join([email.endereco for email in arquivo.emails])
    return render_template(
        "admin/assistencia/atestados/editar.html",
        arquivo=arquivo,
        emails=emails,
    )


@atestados_bp.route("/<int:arquivo_id>/remover", methods=["POST"])
@login_required
def remover(arquivo_id: int):
    arquivo = AtestadoArquivo.query.get_or_404(arquivo_id)
    db.session.delete(arquivo)
    db.session.commit()
    flash("Arquivo e emails associados foram removidos com sucesso.", "success")
    return redirect(url_for("atestados_bp.index"))


@atestados_bp.route("/<int:arquivo_id>/converter-pdf", methods=["GET"])
@login_required
def converter_pdf(arquivo_id: int):
    arquivo = AtestadoArquivo.query.get_or_404(arquivo_id)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx_file:
        tmp_docx_file.write(arquivo.conteudo)
        tmp_docx_file_path = tmp_docx_file.name

    tmp_pdf_file_path = tmp_docx_file_path.replace(".docx", ".pdf")

    try:
        _converter_docx_para_pdf(tmp_docx_file_path, tmp_pdf_file_path)

        with open(tmp_pdf_file_path, "rb") as pdf_file:
            pdf_content = pdf_file.read()

        try:
            write_audit_external(
                entity_type="atestados_pdf",
                entity_id=arquivo.id,
                action="convert_pdf",
                message="Conversão de DOCX para PDF.",
                after={"arquivo": arquivo.nome},
            )
        except Exception:
            current_app.logger.exception("Falha ao auditar conversao de PDF.")

        response = make_response(pdf_content)
        response.headers["Content-Type"] = "application/pdf"
        response.headers["Content-Disposition"] = f'attachment; filename="{arquivo.nome}.pdf"'
        return response
    except Exception as exc:
        flash(f"Erro ao converter o arquivo em PDF: {exc}", "danger")
        try:
            write_audit_external(
                entity_type="atestados_pdf",
                entity_id=arquivo.id,
                action="convert_pdf_error",
                message="Falha na conversão de DOCX para PDF.",
                after={"arquivo": arquivo.nome, "erro": str(exc)},
            )
        except Exception:
            current_app.logger.exception("Falha ao auditar erro de conversao de PDF.")
        return redirect(url_for("atestados_bp.index"))
    finally:
        if os.path.exists(tmp_docx_file_path):
            os.remove(tmp_docx_file_path)
        if os.path.exists(tmp_pdf_file_path):
            os.remove(tmp_pdf_file_path)


@atestados_bp.route("/<int:arquivo_id>/baixar-excel", methods=["GET"])
@login_required
def baixar_excel(arquivo_id: int):
    arquivo = AtestadoArquivo.query.get_or_404(arquivo_id)
    nome_arquivo = arquivo.nome
    conteudo_arquivo = arquivo.conteudo

    try:
        write_audit_external(
            entity_type="atestados_excel",
            entity_id=arquivo.id,
            action="download_excel",
            message="Download de arquivo do atestado.",
            after={"arquivo": arquivo.nome},
        )
    except Exception:
        current_app.logger.exception("Falha ao auditar download de atestado.")

    return send_file(
        io.BytesIO(conteudo_arquivo),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=nome_arquivo,
    )


@atestados_bp.route("/enviar", methods=["GET", "POST"])
@login_required
def enviar():
    if request.method == "POST":
        texto_original = (request.form.get("texto_original") or "").strip()
        if not texto_original:
            return jsonify({"error": "Texto original não foi fornecido."}), 400

        novo_texto = (request.form.get("novo_texto") or "").strip()
        if not novo_texto:
            return jsonify({"error": "Novo texto não foi fornecido."}), 400

        data_limite = (request.form.get("data_limite") or "").strip()
        if not data_limite:
            return jsonify({"error": "Data limite não foi fornecida."}), 400

        arquivo_ids = request.form.getlist("arquivo_ids")
        if not arquivo_ids:
            return jsonify({"error": "Nenhum arquivo foi selecionado."}), 400

        task_id = str(uuid.uuid4())
        task = AtestadoTask(
            id=task_id,
            status="running",
            progress=0,
            total=len(arquivo_ids),
            current=0,
        )
        db.session.add(task)
        db.session.commit()

        app = current_app._get_current_object()
        submit_bg_task(app, _processar_envio, app, task_id, novo_texto, data_limite, arquivo_ids, max_retries=1)

        return jsonify({"task_id": task_id}), 202

    arquivos = AtestadoArquivo.query.order_by(AtestadoArquivo.id.asc()).all()
    return render_template(
        "admin/assistencia/atestados/enviar.html",
        arquivos=arquivos,
        nav_counts=_nav_counts(),
        active_tab="envios",
    )


def _normalize_txt(text: str) -> str:
    if not text:
        return ""
    nfkd = unicodedata.normalize("NFKD", text)
    return "".join([c for c in nfkd if not unicodedata.combining(c)]).lower()

def _processar_documento_atestado(document, novo_texto: str) -> bool:
    """
    Substitui flexivelmente 'MES/ANO', 'Mês de Referência: ...' ou qualquer data/mês de referência anterior.
    Percorre o corpo do documento, todas as tabelas, cabeçalhos e rodapés.
    Usa normalização ASCII para garantir imunidade a problemas de encoding de arquivo.
    """
    texto_encontrado = False
    try:
        parts = novo_texto.split("/")
        mes_novo = parts[0].strip().upper()
        ano_novo = parts[1].strip()
    except Exception:
        mes_novo = novo_texto.strip().upper()
        ano_novo = ""

    novo_valor = f"{mes_novo}/{ano_novo}" if ano_novo else mes_novo

    all_paragraphs = list(document.paragraphs)

    def extract_table_paragraphs(tables):
        pars = []
        for t in tables:
            for r in t.rows:
                for c in r.cells:
                    pars.extend(c.paragraphs)
                    if getattr(c, "tables", None):
                        pars.extend(extract_table_paragraphs(c.tables))
        return pars

    all_paragraphs.extend(extract_table_paragraphs(document.tables))

    for section in document.sections:
        if getattr(section, "header", None):
            all_paragraphs.extend(section.header.paragraphs)
            all_paragraphs.extend(extract_table_paragraphs(section.header.tables))
        if getattr(section, "footer", None):
            all_paragraphs.extend(section.footer.paragraphs)
            all_paragraphs.extend(extract_table_paragraphs(section.footer.tables))

    import re
    pattern_mes_ano = re.compile(r"mes\s*/\s*ano", re.IGNORECASE)
    pattern_mes_extenso = re.compile(
        r"\b(janeiro|fevereiro|marco|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)/\d{4}\b",
        re.IGNORECASE
    )

    for p in all_paragraphs:
        txt = p.text
        if not txt:
            continue

        norm = _normalize_txt(txt)
        modified = False

        if "mes/ano" in norm or pattern_mes_ano.search(norm):
            if "MES/ANO" in txt or "MÊS/ANO" in txt or "Mes/Ano" in txt:
                txt = txt.replace("MES/ANO", novo_valor).replace("MÊS/ANO", novo_valor).replace("Mes/Ano", novo_valor)
            else:
                txt = re.sub(r"M[EÊeê]S\s*/\s*ANO", novo_valor, txt, flags=re.IGNORECASE)
            modified = True
            texto_encontrado = True
        elif "mes de referencia:" in norm:
            if ":" in txt:
                prefix, _, _ = txt.partition(":")
                txt = f"{prefix}:  {novo_valor}"
            else:
                txt = f"{txt}  {novo_valor}"
            modified = True
            texto_encontrado = True
        elif pattern_mes_extenso.search(norm):
            txt = re.sub(
                r"\b(JANEIRO|FEVEREIRO|MARÇO|MARCO|ABRIL|MAIO|JUNHO|JULHO|AGOSTO|SETEMBRO|OUTUBRO|NOVEMBRO|DEZEMBRO)/\d{4}\b",
                novo_valor,
                txt,
                flags=re.IGNORECASE
            )
            modified = True
            texto_encontrado = True

        if modified:
            p.text = txt
            if p.runs:
                try:
                    _set_highlight_color(p.runs[0], "yellow")
                except Exception:
                    pass

    has_periodo = any("periodo de referencia" in _normalize_txt(p.text) for p in all_paragraphs)
    if has_periodo and ano_novo and ano_novo.isdigit():
        try:
            num_mes = MESES_MAP.get(mes_novo, 1)
            ultimo_dia_mes = calendar.monthrange(int(ano_novo), num_mes)[1]
            periodo_referencia = f"01/{num_mes:02d}/{ano_novo} até {ultimo_dia_mes:02d}/{num_mes:02d}/{ano_novo}"
            for p in all_paragraphs:
                if "periodo de referencia" in _normalize_txt(p.text):
                    p.text = f"Período de Referência: {periodo_referencia}"
                    if p.runs:
                        try:
                            _set_highlight_color(p.runs[0], "yellow")
                        except Exception:
                            pass
        except Exception as ex:
            current_app.logger.warning("Erro ao calcular período de referência: %s", ex)

    if texto_encontrado:
        return True

    for p in all_paragraphs:
        norm_low = _normalize_txt(p.text)
        if "atesto" in norm_low or "referencia" in norm_low or "servico" in norm_low:
            return True

    return False


def _processar_envio(app, task_id: str, novo_texto: str, data_limite: str, arquivo_ids: list[str]) -> None:
    with app.app_context():
        task = AtestadoTask.query.get(task_id)
        if not task:
            current_app.logger.error("Tarefa %s nao encontrada.", task_id)
            return

        total_files = len(arquivo_ids)
        task.total = total_files
        task.current = 0
        db.session.commit()

        try:
            mes_novo, ano_novo = novo_texto.split("/")
        except ValueError:
            task.status = "failed"
            task.error = "Formato inválido. Use MÊS/ANO."
            db.session.commit()
            return

        mes_novo = mes_novo.strip().upper()
        ano_novo = ano_novo.strip()

        failed_count = 0
        for arquivo_id in arquivo_ids:
            task = AtestadoTask.query.get(task_id)
            if task and task.status == "cancelled":
                current_app.logger.info("Tarefa %s cancelada.", task_id)
                break
            try:
                arquivo_id_int = int(arquivo_id)
            except (TypeError, ValueError):
                arquivo_id_int = None

            arquivo = AtestadoArquivo.query.get(arquivo_id_int) if arquivo_id_int else None
            if not arquivo:
                failed_count += 1
                current_app.logger.error("Arquivo com ID %s não encontrado.", arquivo_id)
                continue

            if task:
                task.current += 1
                task.progress = int((task.current / task.total) * 100)
                db.session.commit()

            tmp_docx_file_path = ""
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx_file:
                    tmp_docx_file.write(arquivo.conteudo)
                    tmp_docx_file_path = tmp_docx_file.name

                document = Document(tmp_docx_file_path)

                texto_encontrado = _processar_documento_atestado(document, novo_texto)

                if not texto_encontrado:
                    failed_count += 1
                    mensagem_erro = f"Texto 'MES/ANO' não encontrado no documento {arquivo.nome}."
                    db.session.add(AtestadoLogEnvio(
                        arquivo_id=arquivo.id,
                        status="error",
                        mensagem=mensagem_erro,
                        data_envio=datetime.now(),
                    ))
                    db.session.commit()
                    continue

                modified_docx_path = tmp_docx_file_path.replace(".docx", "_modified.docx")
                document.save(modified_docx_path)

                _converter_docx_para_pdf(modified_docx_path, modified_docx_path.replace(".docx", ".pdf"))

                status, message = _enviar_email_com_anexo(
                    arquivo,
                    modified_docx_path.replace(".docx", ".pdf"),
                    mes_novo,
                    ano_novo,
                    data_limite,
                )
                if status == "success":
                    emails_enviados = ", ".join([email.endereco for email in arquivo.emails])
                    mensagem_sucesso = (
                        f"Email enviado com sucesso para os destinatarios do arquivo {arquivo.nome}: "
                        f"{emails_enviados}."
                    )
                    db.session.add(AtestadoLogEnvio(
                        arquivo_id=arquivo.id,
                        status="success",
                        mensagem=mensagem_sucesso,
                        data_envio=datetime.now(),
                    ))
                else:
                    failed_count += 1
                    mensagem_erro = f"Erro ao enviar email para {arquivo.nome}: {message}"
                    db.session.add(AtestadoLogEnvio(
                        arquivo_id=arquivo.id,
                        status="error",
                        mensagem=mensagem_erro,
                        data_envio=datetime.now(),
                    ))
                db.session.commit()

                for path in (
                    tmp_docx_file_path,
                    modified_docx_path,
                    modified_docx_path.replace(".docx", ".pdf"),
                ):
                    if os.path.exists(path):
                        os.remove(path)

            except Exception as exc:
                failed_count += 1
                current_app.logger.error("Erro ao processar o arquivo %s: %s", arquivo.nome, exc)
                db.session.add(AtestadoLogEnvio(
                    arquivo_id=arquivo.id,
                    status="error",
                    mensagem=f"Erro ao processar o arquivo {arquivo.nome}: {exc}",
                    data_envio=datetime.now(),
                ))
                db.session.commit()
                continue
            finally:
                if tmp_docx_file_path and os.path.exists(tmp_docx_file_path):
                    try:
                        os.remove(tmp_docx_file_path)
                    except Exception:
                        pass

        else:
            if task:
                task.status = "completed"
                task.progress = 100
                if failed_count > 0:
                    task.error = f"Envio finalizado. {failed_count} arquivo(s) falhou/falharam ao processar ou enviar. Consulte o Histórico para ver mais detalhes."
                db.session.commit()




@atestados_bp.route("/task-status/<task_id>", methods=["GET"])
@login_required
def task_status(task_id: str):
    task = AtestadoTask.query.get(task_id)
    if task:
        return jsonify({"status": task.status, "progress": task.progress, "error": task.error})
    return jsonify({"error": "Tarefa não encontrada."}), 404



@atestados_bp.route("/cancel-task/<task_id>", methods=["POST"])
@login_required
def cancel_task(task_id: str):
    task = AtestadoTask.query.get(task_id)
    if task:
        task.status = "cancelled"
        db.session.commit()
        return jsonify({"message": "Tarefa cancelada com sucesso."})
    return jsonify({"error": "Tarefa não encontrada."}), 404


@atestados_bp.route("/historico", methods=["GET"])
@login_required
def historico_envios():
    page = request.args.get("page", 1, type=int)
    data_inicio = request.args.get("data_inicio", "").strip()
    hora_inicio = request.args.get("hora_inicio", "00:00").strip()
    data_fim = request.args.get("data_fim", "").strip()
    hora_fim = request.args.get("hora_fim", "23:59").strip()
    status = request.args.get("status", "").strip().lower()

    query = AtestadoLogEnvio.query.join(AtestadoLogEnvio.arquivo)

    try:
        if data_inicio:
            dt_inicio = datetime.strptime(f"{data_inicio} {hora_inicio}", "%Y-%m-%d %H:%M")
            query = query.filter(AtestadoLogEnvio.data_envio >= dt_inicio)
        if data_fim:
            dt_fim = datetime.strptime(f"{data_fim} {hora_fim}", "%Y-%m-%d %H:%M")
            dt_fim = dt_fim.replace(second=59, microsecond=999999)
            query = query.filter(AtestadoLogEnvio.data_envio <= dt_fim)
    except ValueError:
        flash("Data ou hora inválida.", "warning")

    if status in ["success", "error"]:
        query = query.filter(AtestadoLogEnvio.status == status)

    logs = (
        query.options(joinedload(AtestadoLogEnvio.arquivo).joinedload(AtestadoArquivo.emails))
        .order_by(AtestadoLogEnvio.data_envio.desc())
        .paginate(page=page, per_page=10)
    )

    total = query.count()
    success_total = query.filter(AtestadoLogEnvio.status == "success").count()
    error_total = query.filter(AtestadoLogEnvio.status == "error").count()
    stats = [
        {"label": "Registros", "value": total},
        {"label": "Sucesso", "value": success_total},
        {"label": "Erro", "value": error_total},
    ]

    return render_template(
        "admin/assistencia/atestados/historico_envios.html",
        logs=logs,
        data_inicio=data_inicio,
        hora_inicio=hora_inicio,
        data_fim=data_fim,
        hora_fim=hora_fim,
        status=status,
        stats=stats,
        nav_counts=_nav_counts(),
        active_tab="historico",
    )


@atestados_bp.route("/historico/pdf", methods=["GET"])
@login_required
def historico_envios_pdf():
    from weasyprint import HTML

    data_inicio = request.args.get("data_inicio", "").strip()
    hora_inicio = request.args.get("hora_inicio", "00:00").strip()
    data_fim = request.args.get("data_fim", "").strip()
    hora_fim = request.args.get("hora_fim", "23:59").strip()
    status = request.args.get("status", "").strip().lower()

    query = AtestadoLogEnvio.query.join(AtestadoLogEnvio.arquivo)

    try:
        if data_inicio:
            dt_inicio = datetime.strptime(f"{data_inicio} {hora_inicio}", "%Y-%m-%d %H:%M")
            query = query.filter(AtestadoLogEnvio.data_envio >= dt_inicio)
        if data_fim:
            dt_fim = datetime.strptime(f"{data_fim} {hora_fim}", "%Y-%m-%d %H:%M")
            dt_fim = dt_fim.replace(second=59, microsecond=999999)
            query = query.filter(AtestadoLogEnvio.data_envio <= dt_fim)
    except ValueError:
        flash("Data ou hora inválida para exportação em PDF.", "warning")

    if status in ["success", "error"]:
        query = query.filter(AtestadoLogEnvio.status == status)

    logs = query.options(joinedload(AtestadoLogEnvio.arquivo).joinedload(AtestadoArquivo.emails)).all()

    html = render_template(
        "admin/assistencia/atestados/historico_envios_pdf.html",
        logs=logs,
        data_inicio=data_inicio,
        hora_inicio=hora_inicio,
        data_fim=data_fim,
        hora_fim=hora_fim,
        status=status,
    )

    pdf_bin = io.BytesIO()
    HTML(string=html, base_url=request.root_url).write_pdf(pdf_bin)
    pdf_bin.seek(0)

    filename = f"historico_envios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return send_file(pdf_bin, mimetype="application/pdf", as_attachment=True, download_name=filename)
